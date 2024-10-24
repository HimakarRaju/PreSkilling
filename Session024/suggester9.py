import os
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import numpy as np
import sympy as sp
from sympy import Eq, symbols
from scipy.optimize import linprog
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.metrics import classification_report
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from docx import Document
from docx.shared import Inches


def load_dataset(file_path):
    """Load the dataset from a CSV or Excel file."""
    try:
        if file_path.endswith(".csv"):
            return pd.read_csv(file_path)
        elif file_path.endswith(".xlsx"):
            return pd.read_excel(file_path)
        else:
            raise ValueError(f"File format not supported: {file_path}")
    except Exception as e:
        print(f"Error loading dataset: {e}")
        raise


def detect_sequential_columns(df):
    """Detect columns that contain sequential numbers (e.g., ID or serial number)."""
    sequential_columns = [
        col for col in df.columns if df[col].diff().dropna().eq(1).all()
    ]
    return sequential_columns


def generate_report(df, file_path):
    """Generate a comprehensive Word document report for the dataset."""
    document = Document()
    dataset_name = os.path.basename(file_path).split(".")[0]

    # Page 1: Dataset Information
    document.add_heading(dataset_name, level=1)
    document.add_paragraph(f"Dataset Name: {dataset_name}")
    document.add_paragraph(f"Shape: {df.shape[0]} rows, {df.shape[1]} columns")

    # Missing values table
    document.add_heading("Missing Values in Dataset:", level=2)
    missing_values = df.isnull().sum()
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Column Name"
    hdr_cells[1].text = "Missing Values"

    for col, missing in missing_values.items():
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = str(missing)

    # Descriptive Statistics Table
    document.add_heading("Dataset Description:", level=2)
    describe_data = df.describe().transpose()  # Transpose for better formatting
    table = document.add_table(rows=1, cols=len(describe_data.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(describe_data.columns):
        hdr_cells[i].text = col

    for index, row in describe_data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        for i, value in enumerate(row):
            row_cells[i + 1].text = str(value)

    # Correlation Plot
    document.add_heading("Correlation Plot:", level=2)
    plot_correlation_matrix(df, dataset_name)  # Generate and save the plot
    document.add_picture(f"{dataset_name}_corr.png", width=Inches(5.0))

    # Suggested optimization and models
    document.add_heading("Suggested Optimization Techniques and Models:", level=2)
    optimization_report = suggest_optimization(df)
    document.add_paragraph(optimization_report)

    # Save the report as a Word document
    doc_output = f"{dataset_name}_report.docx"
    document.save(doc_output)
    print(f"Report saved as {doc_output}")


# New function to generate and solve useful equations
def generate_useful_equations(df, target_variable):
    """Generate useful equations based on numeric data and solve them."""
    report = "Generated Useful Equations:\n"

    # Get numeric columns
    numeric_columns = df.select_dtypes(include=["float64", "int64"]).columns.tolist()

    if target_variable not in numeric_columns:
        return "Target variable not found in numeric columns."

    feature_columns = numeric_columns.copy()
    feature_columns.remove(target_variable)

    # Create symbolic variables
    symbols_dict = {col: symbols(col) for col in feature_columns + [target_variable]}

    # Example: creating a simple linear equation based on correlation or regression coefficients
    X = df[feature_columns]
    y = df[target_variable]

    # Perform linear regression to get coefficients for the equation
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42
    )
    model = LinearRegression()
    model.fit(X_train, y_train)

    # Generate equation based on regression coefficients
    equation = symbols_dict[target_variable] - sum(
        coef * symbols_dict[col] for coef, col in zip(model.coef_, feature_columns)
    )

    # Symbolic equation for the target variable
    eq = Eq(equation, 0)
    report += f"\nEquation for {target_variable}: {sp.pretty(eq)}\n"

    # Solve for each variable (symbolically, solving for one variable at a time)
    for col in feature_columns:
        solution = sp.solve(eq, symbols_dict[col])
        report += f"\nSolving for {col}: {solution}\n"

    return report


def plot_correlation_matrix(df, dataset_name):
    """Generate and save a correlation matrix plot as PNG."""
    plt.figure(figsize=(10, 6))
    corr = df.corr()
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm", square=True)
    plt.title(f"Correlation Matrix of {dataset_name}")
    plt.savefig(f"{dataset_name}_corr.png")
    plt.close()


def suggest_optimization(df):
    """Analyze the dataset, suggest relevant optimization techniques and suitable models."""
    report = "\n\n\n\nSuggested Optimization Techniques and Models:\n"

    # Determine model suitability based on column types
    numeric_columns = df.select_dtypes(include=["float64", "int64"]).columns.tolist()
    categorical_columns = df.select_dtypes(include=["object"]).columns.tolist()

    if numeric_columns:
        report += "\nSuggested Optimization Techniques:\n"

        if len(numeric_columns) > 1:
            # Calculate the correlation matrix for numeric columns
            corr_matrix = df[numeric_columns].corr()

            # Try to find the most correlated variable with the highest correlation
            filtered_corr = corr_matrix.dropna(axis=1, how="all").dropna(
                axis=0, how="all"
            )

            if not filtered_corr.empty:
                target_variable = filtered_corr.mean().idxmax()
                report += f"   - Suggested target variable: {target_variable} (based on highest correlation).\n"
            else:
                target_variable = numeric_columns[
                    -1
                ]  # Fallback if no correlation found
                report += "   - No valid correlations found. Falling back to the last numeric column as the target.\n"
        else:
            target_variable = numeric_columns[
                0
            ]  # Only one numeric column, so use it as the target
            report += f"   - Only one numeric column found. Setting {target_variable} as the target variable.\n"

        # Allow the user to override the chosen target variable
        common_targets = {"target", "label", "y", "price", "cost", "estimate"}
        if target_variable.lower() not in common_targets:
            report += f"Suggested target variable: {target_variable}. Is this acceptable? (y/n)\n"
            print(
                f"Suggested target variable: {target_variable}. Is this acceptable? (y/n)"
            )  # Inform the user
            user_input = input().strip().lower()
            if user_input != "y":
                report += f"The available numeric columns are: {numeric_columns}\n"
                print(
                    f"The available numeric columns are: {numeric_columns}"
                )  # Inform the user about available columns
                target_variable = input(
                    "Please specify your desired target variable: "
                ).strip()  # Add prompt text
                common_targets.add(target_variable.lower())

        # Generate and solve useful equations based on regression first
        equation_report = generate_useful_equations(df, target_variable)
        report += "\nEquation Generation Result:\n" + equation_report

        # Now proceed with linear programming
        if (
            len(numeric_columns) >= 2
        ):  # Need at least 2 numeric columns (1 target + 1 constraint)
            constraint_columns = numeric_columns[
                :-1
            ]  # Use all other numeric columns as constraints
            optimization_result = run_linear_programming(
                df, constraint_columns, target_variable
            )
            report += "\nLinear Programming Result:\n" + optimization_result
        else:
            report += "   - Not enough numeric columns for Linear Programming.\n"

        # Suggest regression
        if len(numeric_columns) > 1:
            report += "   \n- Regression Analysis: Use when predicting a continuous outcome based on one or more predictors.\n"
            regression_result = run_regression(
                df, target_variable
            )  # Pass the selected target variable
            report += "\nRegression Analysis Result:\n" + regression_result

        # Suggest classification models
        if categorical_columns:
            report += "   - Classification Models:\n"
            report += "     \n* Logistic Regression: Suitable for binary/multiclass outcomes.\n"
            classification_report_str = run_classification(
                df, target_variable
            )  # Use the target variable
            report += "\nLogistic Regression Report:\n" + classification_report_str

            # Decision Tree
            report += "     \n* Decision Tree Classifier: Can handle both binary and multiclass tasks.\n"
            tree_report = run_decision_tree(
                df, target_variable
            )  # Use the target variable
            report += "\nDecision Tree Classifier Report:\n" + tree_report

            # Random Forest
            report += "     * Random Forest Classifier: A robust ensemble method.\n"
            forest_report = run_random_forest(
                df, target_variable
            )  # Use the target variable
            report += "\nRandom Forest Classifier Report:\n" + forest_report

    return report


def split_data(df, target_variable):
    """Helper to split data into training and testing sets."""
    feature_columns = df.select_dtypes(include=["float64", "int64"]).columns.tolist()
    if target_variable not in feature_columns:
        return None, None, None, None
    feature_columns.remove(target_variable)
    X = df[feature_columns]
    y = df[target_variable]
    return train_test_split(X, y, test_size=0.3, random_state=42)


def run_regression(df, target_variable):
    X_train, X_test, y_train, y_test = split_data(df, target_variable)
    if X_train is None:
        return "Target variable not found in feature columns."
    model = LinearRegression()
    model.fit(X_train, y_train)
    return f"Regression Coefficients: {model.coef_}, Intercept: {model.intercept_}"


def run_classification(df, target_variable):
    """Run a simple logistic regression classification."""
    report = ""

    feature_columns = df.select_dtypes(include=["float64", "int64"]).columns.tolist()
    if target_variable not in feature_columns:
        return "Target variable not found in feature columns."

    X = df[feature_columns].drop(columns=[target_variable])
    y = df[target_variable]

    # Split the data into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42
    )

    # Create a logistic regression model
    model = LogisticRegression(max_iter=200)
    model.fit(X_train, y_train)

    # Make predictions
    y_pred = model.predict(X_test)

    # Generate and return classification report
    return classification_report(y_test, y_pred)


def run_decision_tree(df, target_variable):
    """Run a Decision Tree Classifier."""
    feature_columns = df.select_dtypes(include=["float64", "int64"]).columns.tolist()
    if target_variable not in feature_columns:
        return "Target variable not found in feature columns."

    X = df[feature_columns].drop(columns=[target_variable])
    y = df[target_variable]

    # Split the data into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42
    )

    # Create a Decision Tree model
    model = DecisionTreeClassifier()
    model.fit(X_train, y_train)

    # Make predictions
    y_pred = model.predict(X_test)

    # Generate and return classification report
    return classification_report(y_test, y_pred)


def run_random_forest(df, target_variable):
    """Run a Random Forest Classifier."""
    feature_columns = df.select_dtypes(include=["float64", "int64"]).columns.tolist()
    if target_variable not in feature_columns:
        return "Target variable not found in feature columns."

    X = df[feature_columns].drop(columns=[target_variable])
    y = df[target_variable]

    # Split the data into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42
    )

    # Create a Random Forest model
    model = RandomForestClassifier()
    model.fit(X_train, y_train)

    # Make predictions
    y_pred = model.predict(X_test)

    # Generate and return classification report
    return classification_report(y_test, y_pred)


def run_linear_programming(df, constraint_columns, target_variable):
    """Run linear programming on the dataset."""
    try:
        c = -1 * df[target_variable].values  # Minimization problem
        A = df[constraint_columns].values
        b = np.zeros(A.shape[0])
        bounds = [(0, None) for _ in range(A.shape[1])]
        res = linprog(c, A_ub=A, b_ub=b, bounds=bounds, method="highs")
        return (
            f"Optimal solution found: {res.x} with target value: {-res.fun}"
            if res.success
            else f"Failed: {res.message}"
        )
    except Exception as e:
        return f"Linear programming failed: {str(e)}"


def main():
    file_path = input("Enter the path to your CSV or Excel file: ")
    df = load_dataset(file_path)

    print("Loading dataset...")

    # Detect and drop sequential columns
    sequential_columns = detect_sequential_columns(df)
    if sequential_columns:
        df = df.drop(columns=sequential_columns)
        print(f"Dropped sequential columns: {sequential_columns}")

    # Generate a comprehensive report
    print("Generating report...")  # Add this prompt
    generate_report(df, file_path)


if __name__ == "__main__":
    main()
