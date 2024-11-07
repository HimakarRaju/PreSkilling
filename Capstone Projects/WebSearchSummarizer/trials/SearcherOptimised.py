import time
import re
import gc

import requests
from bs4 import BeautifulSoup
from docx import Document
import torch
from langchain.chains import LLMChain  # Import from langchain.chains
from langchain_community.llms.huggingface_pipeline import HuggingFacePipeline  # Import from langchain_community
from langchain_core.prompts import PromptTemplate  # Import from langchain_core
from tqdm import tqdm
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline


def measure_time(func):
    """Decorator to measure execution time of a function."""

    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        print(f"{func.__name__} execution time: {end_time - start_time:.2f} seconds")
        return result

    return wrapper


# --- Optimized Model Loading ---
@measure_time
def load_model():
    model_dir = r"C:\gemma-transformers-1.1-2b-it-v1"
    tokenizer = AutoTokenizer.from_pretrained(model_dir)
    model = AutoModelForCausalLM.from_pretrained(model_dir)

    # Dynamic quantization for CPU efficiency
    model = torch.quantization.quantize_dynamic(model, {torch.nn.Linear}, dtype=torch.qint8)

    if torch.cuda.is_available():
        device = torch.device("cuda")
        print(f"Using CUDA: {torch.version.cuda}, Torch: {torch.__version__}")
    else:
        device = torch.device("cpu")
        print("Using CPU")

    model.to(device)
    print("Model loaded")
    return model, tokenizer, device


# --- LangChain Integration ---
def generate_response(prompt, model, tokenizer, device):
    llm = HuggingFacePipeline(pipeline=pipeline(
        "text-generation", model=model, tokenizer=tokenizer, device=device
    ))

    # Simplified prompt templates with LangChain
    prompt_template = PromptTemplate(
        input_variables=["context", "instructions"],
        template="You are a helpful AI assistant.\nContext: {context}\nInstructions: {instructions}"
    )

    llm_chain = LLMChain(llm=llm, prompt=prompt_template)

    # Generate with LangChain for better prompt management
    response = llm_chain.run({"context": prompt["context"], "instructions": prompt["instructions"]})
    return response


# ---  Web Scraping and Prompting ---
# ... (extract_text_from_url and scrape_google_search remain the same)

@measure_time
def get_user_prompt_choice():
    """Displays prompt options and gets the user's choice."""
    print("Choose a prompt style:")
    print("1. Emphasize detail and analysis")
    print("2. Guide exploration of related topics")
    print("3. Encourage a structured response")
    print("4. Combine approaches")
    while True:
        try:
            choice = int(input("Enter your choice (1-4): "))
            if 1 <= choice <= 4:
                return choice
            else:
                print("Invalid choice. Please enter a number between 1 and 4.")
        except ValueError:
            print("Invalid input. Please enter a number.")


@measure_time
def generate_prompt(choice, search_results):
    """Generates the prompt based on the user's choice."""
    if choice == 1:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you please provide a detailed explanation of this topic, including:
        * A comprehensive summary of the key points.
        * An in-depth analysis of the relevant concepts and their implications.
        * Exploration of related topics and how they connect to the main subject.
        Please be thorough and elaborate on the details, providing a rich and informative response. 
        """
    elif choice == 2:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you delve deeper into this subject, exploring the following:
        * How are these topics connected to the main concept?
        * Are there any interesting relationships or contradictions between them?
        * What are some broader implications or applications of these ideas?
        Please provide a comprehensive and insightful analysis. 
        """
    elif choice == 3:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you please elaborate on this topic with a well-structured response that includes:
        * **Introduction:** A brief overview of the subject.
        * **Tabulated Information:** Tabulate the data which can be tabulated
        * **Key Points:** A detailed summary of the most important aspects.
        * **Analysis:** An in-depth examination of the underlying concepts and their significance.
        * **Related Topics:** Exploration of connected ideas and their relevance.
        * **Conclusion:** A concise summary of the key takeaways.
        Please provide a comprehensive and informative response.
        """
    else:  # choice == 4
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you please provide a detailed and comprehensive analysis of this subject, making sure to:
        * Summarize the key points in a clear and concise manner.
        * Explore the connections between the topics.
        * Provide in-depth explanations and examples to illustrate the concepts.
        * Analyze the implications and applications of these ideas.
        * Tabulate the data which can be tabulated
        I'm looking for a thorough and insightful response that expands on the information provided. 
        """
    return prompt


@measure_time
def create_docx(query, gemma_response, docx_filename):
    """Creates a DOCX with the query as the title and formatted response."""
    doc = Document()
    doc.add_heading(query, level=1)

    # Split the response by "**" and add each section as a heading or paragraph
    sections = gemma_response.split("**")
    for i, section in enumerate(sections):
        if section.strip():
            if i > 0:  # Skip the first section (which might be the title)
                doc.add_heading(section.strip(), level=2)  # Use level 2 for subheadings
            else:
                doc.add_paragraph(section.strip())

    # Delete the first two pages (unwanted info) - using OxmlElement
    for _ in range(2):
        if len(doc.element.body) > 0:
            for p in doc.element.body:
                p.getparent().remove(p)

    doc.save(docx_filename)


@measure_time
def scrape_google_search(query, num_results=10):
    """Scrapes Google Search results for a given query."""
    url = f"https://www.google.com/search?q={query}"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/58.0.3029.110 Safari/537.3"
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        print(f"Response status code: {response.status_code}")

        soup = BeautifulSoup(response.content, "html.parser")
        search_results = []
        result_divs = soup.find_all("div", class_="Gx5Zad")  # This might need to be adjusted if Google changes its HTML
        print(f"Found {len(result_divs)} result divs")

        for i in tqdm(range(min(num_results, len(result_divs))), desc="Processing results"):
            result = result_divs[i]
            try:
                title_elem = result.find("h3")
                title = title_elem.text if title_elem else ""
                link_elem = result.find("a")
                link = link_elem["href"] if link_elem else ""
                snippet_elem = result.find("div", class_="VwiC3b")  # This might need to be adjusted
                snippet = snippet_elem.text if snippet_elem else ""
                text = f"{title}\n{link}\n{snippet}"
                if text:
                    search_results.append(text)
            except Exception as e:
                print(f"Error extracting text from result: {e}")
        return search_results
    except requests.exceptions.RequestException as e:
        print(f"Error during web scraping: {e}")
        return []


@measure_time
def main():
    model, tokenizer, device = load_model()
    while True:
        question = input("Enter your Question (or type 'exit' to quit): ")
        if question.lower() == 'exit':
            break

        search_results = scrape_google_search(question)
        if search_results:
            choice = get_user_prompt_choice()
            prompt = generate_prompt(choice, search_results)

            # Efficient response generation
            with torch.no_grad():  # Reduce memory usage
                gemma_response = generate_response(prompt, model, tokenizer, device)

                if prompt in gemma_response:
                    adjusted_response = gemma_response.replace(prompt, "")

                docx_filename = f"ai_response_{int(time.time())}.docx"
                create_docx(question, adjusted_response, docx_filename)
                print(adjusted_response)
            gc.collect()  # Garbage collection
            torch.cuda.empty_cache()  # Clear CUDA cache
        else:
            print("No answer found.")


if __name__ == "__main__":
    main()