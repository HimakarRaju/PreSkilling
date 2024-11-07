import os
os.environ["TF_ENABLE_ONEDNN_OPTS"]="0"
import time
import logging
import requests
import torch
from bs4 import BeautifulSoup
from docx import Document
from tqdm import tqdm
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline, BitsAndBytesConfig

# Set up logging
logging.basicConfig(level=logging.DEBUG)  # Capture all levels of logs
logger = logging.getLogger(__name__)


def measure_time(func):
    """Decorator to measure execution time of a function."""

    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        logger.debug(f"{func.__name__} execution time: {end_time - start_time:.2f} seconds")
        return result

    return wrapper


@measure_time
def load_model():
    """Loads the model and tokenizer once."""
    quantization_config = BitsAndBytesConfig(load_in_4bit=True, bnb_4bit_compute_dtype="float16")
    # Load your Gemma model and tokenizer
    model_dir = "google/gemma-1.1-2b-it"  # this model is from huggingface.com
    model = AutoModelForCausalLM.from_pretrained(model_dir)
    tokenizer = AutoTokenizer.from_pretrained(model_dir, quantization_config=quantization_config)

    # Check GPU availability and use it if possible
    if torch.cuda.is_available():
        logger.info(f"Torch Version: {torch.__version__}")
        logger.info(f"Torch CUDA Version: {torch.version.cuda}")
        device = torch.device("cuda")
    else:
        logger.warning("CUDA not found, using CPU")
        device = torch.device("cpu")
    logger.info(f"Using {device}")
    model.to(device)
    logger.info("Model loaded successfully!")
    return model, tokenizer, device


@measure_time
def generate_response(prompt, generator):
    """Generates a response using the provided generator (text-generation pipeline)."""
    try:
        with tqdm(total=1, desc="Generating response") as pbar:
            outputs = generator(prompt, max_new_tokens=300)
            pbar.update(1)

        answer = outputs[0]["generated_text"]

        if answer:
            cleaned_answer = answer.replace(prompt, "")

    except Exception as e:
        logger.error(f"Error generating response: {e}")
        cleaned_answer = "An error occurred during response generation."

    return cleaned_answer


@measure_time
def extract_text_from_url(url):
    """Extracts the main text content from a given URL."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text(separator=" ", strip=True)
        return text
    except requests.exceptions.RequestException as e:
        logger.error(f"Error fetching URL: {e}")
        return ""


@measure_time
def scrape_google_search(query, num_results=10):
    """Scrapes Google Search results for a given query."""
    url = f"https://www.google.com/search?q={query}"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/58.0.3029.110 Safari/537.3"
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        logger.info(f"Response status code: {response.status_code}")

        soup = BeautifulSoup(response.content, "html.parser")
        search_results = []
        result_divs = soup.find_all("div", class_="Gx5Zad")  # This might need to be adjusted if Google changes its HTML
        logger.debug(f"Found {len(result_divs)} result divs")

        for i in tqdm(range(min(num_results, len(result_divs))), desc="Processing results"):
            result = result_divs[i]
            try:
                title_elem = result.find("h3")
                title = title_elem.text if title_elem else ""
                link_elem = result.find("a")
                link = link_elem["href"] if link_elem else ""
                snippet_elem = result.find("div", class_="VwiC3b")  # This might need to be adjusted
                snippet = snippet_elem.text if snippet_elem else ""
                text = f"{title}\n{link}\n{snippet}"
                if text:
                    search_results.append(text)
            except Exception as e:
                logger.error(f"Error extracting text from result: {e}")
        return search_results
    except requests.exceptions.RequestException as e:
        logger.error(f"Error during web scraping: {e}")
        return []


@measure_time
def get_user_prompt_choice():
    """Displays prompt options and gets the user's choice."""
    print("Choose a prompt style:")
    print("1. Emphasize detail and analysis")
    print("2. Guide exploration of related topics")
    print("3. Encourage a structured response")
    print("4. Combine approaches")
    while True:
        try:
            choice = int(input("Enter your choice (1-4): "))
            if 1 <= choice <= 4:
                return choice
            else:
                print("Invalid choice. Please enter a number between 1 and 4.")
        except ValueError:
            print("Invalid input. Please enter a number.")


@measure_time
def generate_prompt(choice, search_results):
    """Generates the prompt based on the user's choice."""
    if choice == 1:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you please provide a detailed explanation of this topic, including:
        * A comprehensive summary of the key points.
        * An in-depth analysis of the relevant concepts and their implications.
        * Exploration of related topics and how they connect to the main subject.
        Please be thorough and elaborate on the details, providing a rich and informative response."""
    elif choice == 2:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you delve deeper into this subject, exploring the following:
        * How are these topics connected to the main concept?
        * Are there any interesting relationships or contradictions between them?
        * What are some broader implications or applications of these ideas?
        Please provide a comprehensive and insightful analysis."""
    elif choice == 3:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you please elaborate on this topic with a well-structured response that includes:
        * **Introduction:** A brief overview of the subject.
        * **Tabulated Information:** Tabulate the data which can be tabulated.
        * **Key Points:** A detailed summary of the most important aspects.
        * **Analysis:** An in-depth examination of the underlying concepts and their significance.
        * **Related Topics:** Exploration of connected ideas and their relevance.
        * **Conclusion:** A concise summary of the key takeaways.
        Please provide a comprehensive and informative response."""
    else:  # choice == 4
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {search_results}
        Can you please provide a detailed and comprehensive analysis of this subject, making sure to:
        * Summarize the key points in a clear and concise manner.
        * Explore the connections between the topics.
        * Provide in-depth explanations and examples to illustrate the concepts.
        * Analyze the implications and applications of these ideas.
        * Tabulate the data which can be tabulated.
        I'm looking for a thorough and insightful response that expands on the information provided."""
    return prompt


@measure_time
def create_docx(query, gemma_response, docx_filename):
    """Creates a DOCX with the query as the title and formatted response."""
    doc = Document()
    doc.add_heading(query, level=1)

    # Split the response by "**" and add each section as a heading or paragraph
    sections = gemma_response.split("**")
    for i, section in enumerate(sections):
        if section.strip():
            if i > 0:  # Skip the first section (which might be the title)
                doc.add_heading(section.strip(), level=2)  # Use level 2 for subheadings
            else:
                doc.add_paragraph(section.strip())

    doc.save(docx_filename)


# Main loop
@measure_time
def main():
    model, tokenizer, device = load_model()
    generator = pipeline("text-generation", model=model, tokenizer=tokenizer)

    while True:
        question = input("Enter your Question (or type 'exit' to quit): ")
        if question.lower() == "exit":
            break

        search_results = scrape_google_search(question)

        if search_results:
            choice = get_user_prompt_choice()
            prompt = generate_prompt(choice, search_results)
            gemma_response = generate_response(prompt, generator)
            docx_filename = f"ai_optimised_response_{int(time.time())}.docx"
            create_docx(question, gemma_response, docx_filename)
            logger.debug(f"Generated Response:\n{gemma_response}")
            logger.debug(f"Response saved in {docx_filename}")
        else:
            logger.warning("No search results found for your query.")


if __name__ == "__main__":
    main()
