import time
import requests
import torch
from bs4 import BeautifulSoup
from docx import Document
from tqdm import tqdm
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline


def measure_time(func):
    """Decorator to measure execution time of a function."""

    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        print(f"{func.__name__} execution time: {end_time - start_time:.2f} seconds")
        return result

    return wrapper


@measure_time
def load_model():
    """Loads the model and tokenizer for text generation."""
    model_dir = "google/gemma-1.1-2b-it"
    tokenizer = AutoTokenizer.from_pretrained(model_dir)
    model = AutoModelForCausalLM.from_pretrained(model_dir)

    # Optimize with GPU, if available
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)

    print(f"Using device: {device}")
    print("Model loaded successfully!")

    return model, tokenizer, device


@measure_time
def generate_response(prompt, model, tokenizer, device):
    """Generates a response for a given prompt using the loaded model."""
    try:
        generator = pipeline(
            "text-generation", model=model, tokenizer=tokenizer, device=device.index if device.type == 'cuda' else -1
        )
        outputs = generator(prompt, max_new_tokens=300)
        answer = outputs[0]["generated_text"]

        # Remove prompt from the response if present
        cleaned_answer = answer.replace(prompt, "").strip()
        return cleaned_answer

    except Exception as e:
        print(f"Error generating response: {e}")
        return "An error occurred during response generation."


@measure_time
def extract_text_from_url(url):
    """Extracts main text content from the given URL."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")

        # Remove script/style elements and get clean text
        for tag in soup(["script", "style"]):
            tag.extract()
        return soup.get_text(separator=" ", strip=True)

    except requests.exceptions.RequestException as e:
        print(f"Error fetching URL: {e}")
        return ""


@measure_time
def scrape_web_search(query, num_results=10):
    """Scrapes Google Search for results related to a query."""
    url = f"https://www.bing.com/search?q={query}"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/58.0.3029.110 Safari/537.3"
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, "html.parser")
        search_results = []

        result_divs = soup.find_all("div", class_="Gx5Zad")

        for result in result_divs[:num_results]:
            try:
                title_elem = result.find("h3")
                title = title_elem.text if title_elem else ""
                link_elem = result.find("a")
                link = link_elem["href"] if link_elem else ""
                snippet_elem = result.find("div", class_="VwiC3b")
                snippet = snippet_elem.text if snippet_elem else ""

                if title and link:
                    search_results.append(f"{title}\n{link}\n{snippet}")
            except Exception as e:
                print(f"Error extracting text from result: {e}")

        return search_results

    except requests.exceptions.RequestException as e:
        print(f"Error during web scraping: {e}")
        return []


def get_user_prompt_choice():
    """Gets the user's choice for prompt generation style."""
    print("Choose a prompt style:")
    print("1. Emphasize detail and analysis")
    print("2. Guide exploration of related topics")
    print("3. Encourage a structured response")
    print("4. Combine approaches")

    while True:
        try:
            choice = int(input("Enter your choice (1-4): "))
            if 1 <= choice <= 4:
                return choice
            print("Invalid choice. Please enter a number between 1 and 4.")
        except ValueError:
            print("Invalid input. Please enter a number.")


def generate_prompt(choice, search_results):
    """Generates a prompt based on user's choice and search results."""
    prompt_templates = {
        1: """You are a helpful and informative AI assistant.
            Based on the following web information: {search_results}
            Provide a detailed explanation including a summary, analysis, and exploration of related topics.""",
        2: """You are a helpful and informative AI assistant.
            Based on the following web information: {search_results}
            Provide insights on topic connections, relationships, and broader applications.""",
        3: """You are a helpful and informative AI assistant.
            Based on the following web information: {search_results}
            Structure your response with an introduction, key points, analysis, related topics, and a conclusion.""",
        4: """You are a helpful and informative AI assistant.
            Based on the following web information: {search_results}
            Summarize key points, explore connections, and provide detailed explanations with tabulated data where applicable."""
    }
    return prompt_templates.get(choice, "").format(search_results=search_results)


@measure_time
def create_docx(query, gemma_response, docx_filename):
    """Creates a DOCX file with the given query and response."""
    doc = Document()
    doc.add_heading(query, level=1)
    doc.add_paragraph(gemma_response.strip())
    doc.save(docx_filename)
    print(f"Document saved as {docx_filename}")


@measure_time
def main():
    model, tokenizer, device = load_model()
    while True:
        question = input("Enter your Question (or type 'exit' to quit): ")
        if question.lower() == "exit":
            break

        search_results = scrape_google_search(question)
        if search_results:
            choice = get_user_prompt_choice()
            prompt = generate_prompt(choice, search_results)
            gemma_response = generate_response(prompt, model, tokenizer, device)
            docx_filename = f"ai_response_{int(time.time())}.docx"
            create_docx(question, gemma_response, docx_filename)
            print(gemma_response)
        else:
            print("No relevant information found.")


if __name__ == "__main__":
    main()
