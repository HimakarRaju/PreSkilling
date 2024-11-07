import dearpygui.dearpygui as dpg
import time
import re

import requests
import torch
from bs4 import BeautifulSoup
from docx import Document
from tqdm import tqdm
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline


def measure_time(func):
    """Decorator to measure execution time of a function."""

    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        print(f"{func.__name__} execution time: {end_time - start_time:.2f} seconds")
        return result

    return wrapper


@measure_time
def load_model():
    # Load your Gemma model and tokenizer
    model_dir = r"C:\gemma-transformers-1.1-2b-it-v1"  # Replace with the actual path to your model
    tokenizer = AutoTokenizer.from_pretrained(model_dir)
    model = AutoModelForCausalLM.from_pretrained(model_dir)

    # Check GPU availability and use it if possible
    if torch.cuda.is_available():
        print(f'Torch Version: {torch.__version__}')
        print(f'Torch CUDA Version: {torch.version.cuda}')
        device = torch.device("cuda")
    else:
        print("cuda not found using CPU")
        device = torch.device("cpu")
    print(f"Using {device}")
    model.to(device)
    print("Model loaded successfully!")
    return model, tokenizer, device


@measure_time
def generate_response(prompt, model, tokenizer, device):
    """Generates a response using your Gemma model."""
    try:
        generator = pipeline("text-generation", model=model, tokenizer=tokenizer, device=device)
        with tqdm(total=1, desc="Generating response") as pbar:
            outputs = generator(prompt, max_new_tokens=500)
            pbar.update(1)
            time.sleep(0.01)  # You can adjust or remove this delay if needed

        answer = outputs[0]["generated_text"]

        # Extract the prompt style from the original prompt
        prompt_style_match = re.search(r"\*\*(.+?)\*\*", prompt)
        if prompt_style_match:
            prompt_style = prompt_style_match.group(1)

            # Use regex to remove the chosen prompt from the answer
            answer = re.sub(rf"\*\*{prompt_style}\*\*(.*)", r"\1", answer, flags=re.DOTALL).strip()

    except Exception as e:
        print(f"Error generating response: {e}")
        answer = "An error occurred during response generation."

    return answer


@measure_time
def extract_text_from_url(url):
    """Extracts the main text content from a given URL."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text(separator=' ', strip=True)
        return text
    except requests.exceptions.RequestException as e:
        print(f"Error fetching URL: {e}")
        return ""


@measure_time
def scrape_google_search(query, num_results=10):
    """Scrapes Google Search results for a given query."""
    url = f"https://www.google.com/search?q={query}"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/58.0.3029.110 Safari/537.3"
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        print(f"Response status code: {response.status_code}")

        soup = BeautifulSoup(response.content, "html.parser")
        search_results = []
        result_divs = soup.find_all("div", class_="Gx5Zad")  # This might need to be adjusted if Google changes its HTML
        print(f"Found {len(result_divs)} result divs")

        for i in tqdm(range(min(num_results, len(result_divs))), desc="Processing results"):
            result = result_divs[i]
            try:
                title_elem = result.find("h3")
                title = title_elem.text if title_elem else ""
                link_elem = result.find("a")
                link = link_elem["href"] if link_elem else ""
                snippet_elem = result.find("div", class_="VwiC3b")  # This might need to be adjusted
                snippet = snippet_elem.text if snippet_elem else ""
                text = f"{title}\n{link}\n{snippet}"
                if text:
                    search_results.append(text)
            except Exception as e:
                print(f"Error extracting text from result: {e}")
        return search_results
    except requests.exceptions.RequestException as e:
        print(f"Error during web scraping: {e}")
        return []


@measure_time
def get_user_prompt_choice():
    """Displays prompt options and gets the user's choice."""
    print("Choose a prompt style:")
    print("1. Emphasize detail and analysis")
    print("2. Guide exploration of related topics")
    print("3. Encourage a structured response")
    print("4. Combine approaches")
    while True:
        try:
            choice = int(input("Enter your choice (1-4): "))
            if 1 <= choice <= 4:
                return choice
            else:
                print("Invalid choice. Please enter a number between 1 and 4.")
        except ValueError:
            print("Invalid input. Please enter a number.")


@measure_time
def generate_prompt(choice, search_results):
    """Generates the prompt based on the user's choice."""
    # Convert choice to integer (it might be coming in as a string)
    choice = int(choice)

    if choice == 1:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {', '.join(search_results)}
        Can you please provide a detailed explanation of this topic, including:
        * A comprehensive summary of the key points.
        * An in-depth analysis of the relevant concepts and their implications.
        * Exploration of related topics and how they connect to the main subject.
        * **Tabulated Information:** Tabulate the data which can be tabulated
        Please be thorough and elaborate on the details, providing a rich and informative response. 
        """
    elif choice == 2:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {', '.join(search_results)}
        Can you delve deeper into this subject, exploring the following:
        * How are these topics connected to the main concept?
        * Are there any interesting relationships or contradictions between them?
        * What are some broader implications or applications of these ideas?
        * **Tabulated Information:** Tabulate the data which can be tabulated
        Please provide a comprehensive and insightful analysis. 
        """
    elif choice == 3:
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {', '.join(search_results)}
        Can you please elaborate on this topic with a well-structured response that includes:
        * **Introduction:** A brief overview of the subject.
        * **Tabulated Information:** Tabulate the data which can be tabulated
        * **Key Points:** A detailed summary of the most important aspects.
        * **Analysis:** An in-depth examination of the underlying concepts and their significance.
        * **Related Topics:** Exploration of connected ideas and their relevance.
        * **Conclusion:** A concise summary of the key takeaways.
        Please provide a comprehensive and informative response.
        """
    else:  # choice == 4
        prompt = f"""You are a helpful and informative AI assistant. 
        Here's the information I found from the web: {', '.join(search_results)}
        Can you please provide a detailed and comprehensive analysis of this subject, making sure to:
        * Summarize the key points in a clear and concise manner.
        * Explore the connections between the topics.
        * Provide in-depth explanations and examples to illustrate the concepts.
        * Analyze the implications and applications of these ideas.
        * **Tabulated Information:** Tabulate the data which can be tabulated
        I'm looking for a thorough and insightful response that expands on the information provided. 
        """
    return prompt


@measure_time
def create_docx(query, gemma_response, docx_filename):
    """Creates a DOCX with the query as the title and formatted response."""
    doc = Document()
    doc.add_heading(query, level=1)

    # Split the response by "**" and add each section as a heading or paragraph
    sections = gemma_response.split("**")
    for i, section in enumerate(sections):
        if section.strip():
            if i > 0:  # Skip the first section (which might be the title)
                doc.add_heading(section.strip(), level=2)  # Use level 2 for subheadings
            else:
                doc.add_paragraph(section.strip())

    # Delete the first two pages (unwanted info) - using OxmlElement
    for _ in range(2):
        if len(doc.element.body) > 0:
            for p in doc.element.body:
                p.getparent().remove(p)

    doc.save(docx_filename)


# Load the model
model, tokenizer, device = load_model()


def estimate_response_time(prompt, model, tokenizer, device):
    """Estimates the response time based on a test generation."""
    start_time = time.time()
    try:
        # Generate a short test response (e.g., 10 tokens)
        generator = pipeline("text-generation", model=model, tokenizer=tokenizer, device=device)
        generator(prompt, max_new_tokens=10)
    except Exception as e:
        print(f"Error in estimation: {e}")
        return 30  # Default ETA if estimation fails

    end_time = time.time()
    estimation_time = end_time - start_time

    # Extrapolate to the actual response length (500 tokens)
    actual_time = estimation_time * (500 / 10)
    return int(actual_time)


def ask_question(sender, app_data):
    question = dpg.get_value("input_question")
    dpg.configure_item("ask_button", enabled=False)  # Disable the button
    dpg.set_value("response_text", "Scraping Google...")  # Provide feedback in the main window
    search_results = scrape_google_search(question)

    if search_results:
        choice = dpg.get_value("prompt_choice")
        prompt = generate_prompt(choice, search_results)

        estimated_time = estimate_response_time(prompt, model, tokenizer, device)

        # Show the response window and provide feedback
        dpg.configure_item("response_window", show=True)
        dpg.set_value("response_window_text", f"Generating response... (ETA: ~{estimated_time} seconds)")

        # Add a spinner within the 'response_window'
        with dpg.group(parent="response_window_group", horizontal=True):
            spinner_id = dpg.add_loading_indicator(color=(255, 255, 255))  # Store spinner ID
            dpg.add_text("Generating response...")

        gemma_response = generate_response(prompt, model, tokenizer, device)

        # Remove the spinner after response generation
        dpg.delete_item(spinner_id)
        dpg.set_value("response_window_text", gemma_response)

        docx_filename = f"GUI_responses\\ai_response_{int(time.time())}.docx"
        create_docx(question, gemma_response, docx_filename)

    else:
        dpg.set_value("response_text", "No answer found.")  # Provide feedback in the main window

    dpg.configure_item("ask_button", enabled=True)  # Re-enable the button


# Set up DearPyGUI context
dpg.create_context()

# Define global theme
with dpg.theme() as global_theme:
    with dpg.theme_component(dpg.mvAll):
        dpg.add_theme_color(dpg.mvThemeCol_Text, (255, 255, 255))
        dpg.add_theme_color(dpg.mvThemeCol_WindowBg, 
        (30, 30, 30))
        dpg.add_theme_color(dpg.mvThemeCol_Button, (50, 50, 150))
        dpg.add_theme_color(dpg.mvThemeCol_ButtonHovered, (70, 70, 200))
        dpg.add_theme_color(dpg.mvThemeCol_ButtonActive, (90, 90, 250))
        dpg.add_theme_style(dpg.mvStyleVar_FrameRounding, 5)
        dpg.add_theme_style(dpg.mvStyleVar_WindowPadding, 10, 10)

# Define custom fonts
with dpg.font_registry():
    default_font = dpg.add_font("C:/Windows/Fonts/Arial.ttf", 20)  # Adjust font path and size as needed
    title_font = dpg.add_font("C:/Windows/Fonts/Arial.ttf", 24)  # Larger font for the title

# Apply the theme and font
dpg.bind_theme(global_theme)
dpg.bind_font(default_font)

# Define UI elements
with dpg.window(label="AI Assistant", width=780, height=580, pos=(10, 10), tag="main_window"):
    # ... (Other UI elements in the main window remain the same)

    dpg.add_spacer(height=10)
    dpg.add_text("Response:", color=(255, 255, 255))
    with dpg.group(tag="response_text_group"):
        dpg.add_text("Response will appear in a separate window", tag="response_text", wrap=580)

    # Create a separate window for the response (initially hidden)
    with dpg.window(label="Response", width=700, height=400, pos=(100, 100), show=False, tag="response_window"):
        with dpg.group(tag="response_window_group"):  # Container for the response text
            dpg.add_text("", tag="response_window_text", wrap=680)  # Text for the response

    with dpg.group(horizontal=True):
        dpg.add_input_text(tag="input_question", hint="Enter your question", width=580)
        dpg.add_button(label="Ask", tag="ask_button", callback=ask_question)

    dpg.add_spacer(height=10)
    dpg.add_text("Choose a prompt style:", color=(255, 255, 255))
    dpg.add_radio_button(items=["Detail and Analysis", "Explore Topics", "Structured Response", "Comprehensive"],
                         tag="prompt_choice", default_value=0)

    # dpg.add_spacer(height=10)
    # dpg.add_text("Response:", color=(255, 255, 255))
    # dpg.add_text("", tag="response_text", wrap=580)

# Setup and show viewport
dpg.create_viewport(title="AI Assistant App", width=800, height=600)
dpg.setup_dearpygui()
dpg.show_viewport()
dpg.start_dearpygui()
dpg.destroy_context()
